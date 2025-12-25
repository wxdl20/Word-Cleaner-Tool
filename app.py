import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from tkinter import ttk  # 导入更现代的组件库(用于进度条)
from docx import Document
import re
import os
import threading

def process_single_file(file_path):
    """处理单个文件的逻辑，返回 (是否成功, 信息)"""
    try:
        doc = Document(file_path)
        new_doc = Document()
        has_chinese = False
        
        # 遍历段落
        for para in doc.paragraphs:
            text = para.text.strip()
            # 只要包含中文就保留
            if re.search(r'[\u4e00-\u9fa5]', text):
                new_doc.add_paragraph(para.text)
                has_chinese = True
        
        # 构造新文件名
        dir_name = os.path.dirname(file_path)
        base_name = os.path.basename(file_path)
        name_part = os.path.splitext(base_name)[0]
        new_path = os.path.join(dir_name, f"{name_part}_纯中文版.docx")
        
        new_doc.save(new_path)
        
        if not has_chinese:
            return True, f"完成 (警告:原文档无中文): {base_name}"
        return True, f"成功: {base_name}"
        
    except Exception as e:
        return False, f"失败 {os.path.basename(file_path)}: {str(e)}"

def run_batch_process(files):
    """批量处理线程"""
    total = len(files)
    success_count = 0
    
    # 重置进度条
    progress_bar['maximum'] = total
    progress_bar['value'] = 0
    
    # 启用日志框
    log_text.config(state='normal')
    log_text.delete(1.0, tk.END) # 清空旧日志
    log_text.insert(tk.END, f"--- 开始处理 {total} 个文件 ---\n")
    
    for i, file_path in enumerate(files):
        # 更新状态
        status_label.config(text=f"正在处理 ({i+1}/{total}): {os.path.basename(file_path)}")
        
        # 执行处理
        success, msg = process_single_file(file_path)
        
        # 记录结果
        if success:
            success_count += 1
            log_text.insert(tk.END, f"[√] {msg}\n", 'success')
        else:
            log_text.insert(tk.END, f"[X] {msg}\n", 'error')
        
        # 滚动到底部
        log_text.see(tk.END)
        
        # 更新进度条
        progress_bar['value'] = i + 1
        root.update_idletasks() # 强制刷新界面
        
    # 结束处理
    log_text.insert(tk.END, f"\n--- 处理结束: 成功 {success_count} / 总计 {total} ---\n")
    log_text.config(state='disabled') # 禁止用户编辑日志
    status_label.config(text="全部完成！")
    btn.config(state='normal') # 重新启用按钮
    messagebox.showinfo("完成", f"批量处理完成！\n成功：{success_count}\n失败：{total - success_count}")

def select_files():
    # 允许选择多个文件
    file_paths = filedialog.askopenfilenames(
        title="批量选择Word文档 (可多选)",
        filetypes=[("Word Documents", "*.docx")]
    )
    
    if file_paths:
        # 禁用按钮防止重复点击
        btn.config(state='disabled')
        # 开启线程
        threading.Thread(target=run_batch_process, args=(file_paths,)).start()

# --- 界面布局 ---
root = tk.Tk()
root.title("Word批量去英文工具")
root.geometry("500x450")

# 1. 标题区
header_frame = tk.Frame(root)
header_frame.pack(pady=15)
tk.Label(header_frame, text="Word文档批量处理器", font=("微软雅黑", 16, "bold")).pack()
tk.Label(header_frame, text="支持按住 Ctrl 或 Shift 键一次选择多个文件", fg="#666").pack()

# 2. 按钮区
btn_frame = tk.Frame(root)
btn_frame.pack(pady=10)
btn = tk.Button(btn_frame, text="📂 批量选择并开始", font=("微软雅黑", 11), command=select_files, bg="#0078D7", fg="white", width=20, height=2)
btn.pack()

# 3. 进度条
progress_bar = ttk.Progressbar(root, orient="horizontal", length=400, mode="determinate")
progress_bar.pack(pady=10)

status_label = tk.Label(root, text="等待任务...", fg="blue")
status_label.pack()

# 4. 日志区 (带滚动条)
log_frame = tk.Frame(root)
log_frame.pack(padx=20, pady=10, fill='both', expand=True)

tk.Label(log_frame, text="处理日志:", anchor='w').pack(fill='x')

scrollbar = tk.Scrollbar(log_frame)
scrollbar.pack(side=tk.RIGHT, fill='y')

log_text = tk.Text(log_frame, height=10, state='disabled', yscrollcommand=scrollbar.set, font=("Consolas", 9))
log_text.pack(side=tk.LEFT, fill='both', expand=True)

# 配置日志颜色
log_text.tag_config('success', foreground='green')
log_text.tag_config('error', foreground='red')

scrollbar.config(command=log_text.yview)

root.mainloop()